import datetime
import os
import traceback
from queue import Queue
from threading import Thread

import requests
from lxml import etree
from docx import Document
from docx.shared import Inches
from selenium import webdriver
from selenium.common.exceptions import TimeoutException
import logging


if not os.path.exists('logs'):
    os.mkdir('logs')

today = datetime.datetime.today().date()
today_path = 'logs/%s-runtime.log' % today
with open(today_path, 'a') as f:
    pass


logging.basicConfig(
    filename=today_path,
    filemode='a',
    format='%(asctime)s,%(msecs)d %(name)s %(levelname)s %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S',
    level=logging.INFO
)


class Runner:
    def __init__(self,
                 keywords,
                 url_template,
                 driver_path,
                 paragraph_xpath_list,
                 title_xpath_list,
                 headless=True,
                 links_xpath='//h3/a/@href',
                 debug=True,
                 start_page=1,
                 **kwargs):
        self.keywords = keywords
        self.all_links = []
        self.queue = Queue()
        self.new_links_crawled = []
        self.headless = headless
        self.start_page = start_page
        # self.paragraphs = None
        # self.get_paragraphs()
        self.links_xpath = links_xpath
        self.title_xpath_list = title_xpath_list
        self.paragraph_xpath_list = paragraph_xpath_list
        self.debug = debug
        self.url_template = url_template
        self.test_links = kwargs.get('test_links')
        self.crawled_fn = 'crawled.txt'
        self.kw_not_match_fn = 'kw_not_match.txt'
        self.intermediate_folder_name = 'intermediates'
        self.output_folder_name = 'output'
        self.driver_path = driver_path
        self.init_folders_or_files()

        with open(self.crawled_fn, 'r') as f:
            self.crawled = set([item.rstrip('\n') for item in f.readlines()])

        with open(self.kw_not_match_fn, 'r') as f:
            self.not_matched = set([item.rstrip('\n') for item in f.readlines()])

        self.driver = self.init_driver()
        self.raise_exception = kwargs.get('raise_exception', True)

    def init_folders_or_files(self):
        if not os.path.exists(self.crawled_fn):
            with open(self.crawled_fn, 'w') as fp:
                pass

        if not os.path.exists(self.intermediate_folder_name):
            os.mkdir(self.intermediate_folder_name)

        if not os.path.exists(self.output_folder_name):
            os.mkdir(self.output_folder_name)

        if not os.path.exists(self.crawled_fn):
            with open(self.crawled_fn, 'w') as fp:
                pass

        if not os.path.exists(self.kw_not_match_fn):
            with open(self.kw_not_match_fn, 'w') as fp:
                pass

    def get_paragraphs(self, html):
        for index in range(len(self.paragraph_xpath_list)):
            paragraphs = html.xpath(self.paragraph_xpath_list[index])
            if paragraphs:
                return paragraphs or []
            else:
                continue

        return []

    def save(self, title, url, paragraphs):
        # 新建空白文档
        doc = Document()
        # 新增文档标题
        doc.add_heading(title, 0)
        doc.add_heading('原文链接 ' + url, 3)
        doc.add_picture('%s/%s-start.png' % (self.intermediate_folder_name, title), width=Inches(5))
        doc.add_paragraph()
        doc.add_picture('%s/%s-middle.png' % (self.intermediate_folder_name, title), width=Inches(5))
        doc.add_paragraph()
        doc.add_picture('%s/%s-end.png' % (self.intermediate_folder_name, title), width=Inches(5))
        for p in paragraphs:
            # 创建段落描述
            doc.add_paragraph(p)
        # logging.info('SAVING DOC: %s' % title)

        doc.save('%s/%s.docx' % (self.output_folder_name, title))

    def produce(self):
        logging.info('PRODUCE')
        end_link = 'http:stop'

        if self.debug and self.test_links:
            self.all_links = self.test_links
        else:
            for kw in self.keywords:
                try:
                    self.all_links += self.get_page_item_links(kw, self.start_page)
                except Exception as e:
                    logging.error(str(traceback.format_exc()))

            self.all_links = [l for l in self.all_links if l.startswith('http')]

        self.queue.put((end_link, None))
        self.all_links.append(end_link)

    def init_driver(self):
        logging.info('INIT DRIVER')
        options = webdriver.ChromeOptions()
        if self.headless:
            options.add_argument('--headless')

        driver = webdriver.Chrome(self.driver_path, options=options)
        driver.set_page_load_timeout(10)
        return driver

    def get_page_item_links(self, kw, page, previous_links=None, total_links=None):
        if not total_links:
            total_links = []
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/93.0.4577.82 Safari/537.36'
        }

        url = self.url_template.format(kw=kw, page=page)
        content = requests.get(url, headers=headers).content.decode('utf-8')
        html = etree.HTML(content)
        links = html.xpath(self.links_xpath)
        if not links or links == previous_links:
            return total_links

        for link in links:
            # print('put % into queue' % link)
            self.queue.put((link, kw))

        logging.info('GET PARSED LINKS:\n%s', links)
        total_links += links
        return self.get_page_item_links(kw, page=page + 1, previous_links=links[:], total_links=total_links[:])

    def process_single(self, link, kw):
        if link in self.crawled:
            logging.info('FOUND DUPLICATE LINK: %s, SKIPPING...', link)
            return

        if link in self.not_matched:
            logging.info('FOUND NOT MATCHED LINK: %s, SKIPPING...', link)
            return

        try:
            # d = self.init_driver()
            self.driver.get(link)
            ps = self.driver.page_source
            html = etree.HTML(ps)

            title = html.xpath(self.title_xpath_list[0])
            if not title and len(self.title_xpath_list) > 1:
                title = html.xpath(self.title_xpath_list[1])

            title = title[0]

            if os.path.exists('%s/%s.docx' % (self.output_folder_name, title)):
                logging.info('FOUND DUPLICATE DOC: %s, SKIPPING...', title)
                return

            url = self.driver.current_url
            paragraphs = self.get_paragraphs(html=html)

            if not paragraphs:
                raise ValueError('NO PARAGRAPH: %s...' % url)

            full_text = ''.join(paragraphs)
            if kw not in full_text:
                logging.info('KEYWORD %s NOT FOUND IN %s, SKIPPING...', kw, link)
                with open(self.kw_not_match_fn, 'a') as fp:
                    fp.write(link + '\n')
                return

            logging.info('CRAWLING %s %s', title, link)

            self.driver.get_screenshot_as_file('intermediates/%s-start.png' % title)

            self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight / 2)")
            # time.sleep(0.2)
            self.driver.get_screenshot_as_file('intermediates/%s-middle.png' % title)

            self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight)")
            # time.sleep(0.2)
            self.driver.get_screenshot_as_file('intermediates/%s-end.png' % title)

            self.save(title=title, url=url, paragraphs=paragraphs)
            logging.info('SUCCESSFULLY SAVED %s', title)
            self.new_links_crawled.append(url)

        except TimeoutException:
            return

        except Exception as e:
            if self.debug and self.raise_exception:
                raise e

            logging.error(str(traceback.format_exc()))
            # traceback.print_exc(file=open('err.log', 'a+'))

            with open('failed_to_crawl.txt', 'a') as f:
                f.write(link + '\n')
        else:
            with open(self.crawled_fn, 'a') as f:
                f.write(link + '\n')

    def consume(self):
        while True:
            link, kw = self.queue.get()
            logging.info('RETRIEVED LINK FROM QUEUE: %s', link)
            if link == 'http:stop':
                logging.info('RECEIVED STOP SIGNAL, STOPPED!!!')
                self.driver.close()
                break

            self.process_single(link, kw)
        self.driver.close()

    def run(self):
        t1 = Thread(target=self.produce)
        t2 = Thread(target=self.consume)

        t1.start()
        t2.start()

        t1.join()
        t2.join()
