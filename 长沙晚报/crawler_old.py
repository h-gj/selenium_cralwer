import os
import time
import traceback

import requests
from lxml import etree
from docx import Document
from docx.shared import Inches
from selenium import webdriver


class WordSaver:
    def __init__(self, html, title, url):
        self.html = html
        self.title = title
        self.url = url
        self.paragraphs = None
        self.parse()

    def parse(self):
        self.paragraphs = self.html.xpath('//article[@class="am-article"]//p/text()')
        if not self.paragraphs:
            raise ValueError('Error getting content')
            print('getting content from another selector...')
            self.paragraphs = self.html.xpath('//div[@class="nei"]/ct/p//text()')
            if not self.paragraphs:
                print('failed again to find body: %s...' % self.url)

    def save(self):
        # 新建空白文档
        doc = Document()
        # 新增文档标题
        doc.add_heading(self.title, 0)
        doc.add_heading('原文链接 ' + self.url, 3)
        doc.add_picture('intermediates/%s-start.png' % self.title, width=Inches(5))
        doc.add_paragraph()
        doc.add_picture('intermediates/%s-middle.png' % self.title, width=Inches(5))
        doc.add_paragraph()
        doc.add_picture('intermediates/%s-end.png' % self.title, width=Inches(5))
        for p in self.paragraphs:
            # 创建段落描述
            doc.add_paragraph(p)
        doc.save('output/%s.docx' % self.title)


def get_page_item_links(kw, page=1, total_links=None):
    if not total_links:
        total_links = []
    headers = {
        'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/93.0.4577.82 Safari/537.36'
    }
    url = 'https://so.icswb.com/default.php?mod=search&m=no&syn=no&f=_all&s=s_show_date_DESC&temp=&p=%s&ps=20&site_id=2&range=&search_target=1&search_key=%s&search_column=&search_channel_id=0' % (page, kw)
    content = requests.get(url, headers=headers).content.decode('utf-8')
    html = etree.HTML(content)
    links = html.xpath('//h3/a/@href')
    if not links:
        return total_links

    print('links', links)
    total_links += links
    return get_page_item_links(kw, page=page + 1, total_links=total_links[:])


if __name__ == '__main__':
    keywords = ['望城县', '长沙市望城区', '望城区档案']
    all_links = []
    new_links_crawled = []
    # for kw in keywords:
    #     all_links += get_page_item_links(kw)

    all_links = ['https://www.icswb.com/h/104324/20201230/691731.html', 'https://www.icswb.com/h/149/20201224/690816.html', 'https://www.icswb.com/h/101150/20201223/690769.html', 'https://www.icswb.com/h/150/20201223/690739.html', 'https://www.icswb.com/h/161/20201223/690796.html', 'https://www.icswb.com/h/100104/20201223/690654.html', 'https://www.icswb.com/h/151/20201222/690560.html', 'https://www.icswb.com/h/151/20201221/690357.html', 'https://www.icswb.com/h/100104/20201220/690303.html', 'https://www.icswb.com/h/150/20201219/690166.html', 'https://www.icswb.com/h/160/20201218/690136.html', 'https://www.icswb.com/h/101765/20201218/690030.html', 'https://www.icswb.com/h/100104/20201218/690087.html', 'https://www.icswb.com/h/162/20201217/689868.html', 'https://www.icswb.com/h/162/20201217/689918.html', 'https://www.icswb.com/h/161/20201217/689902.html', 'https://www.icswb.com/h/162/20201217/689836.html', 'https://www.icswb.com/h/101946/20201217/689823.html', 'https://www.icswb.com/h/101321/20201217/689798.html', 'https://www.icswb.com/h/100040/20201216/689675.html']
    # all_links = ['http://www.chinaleifeng.com/content/2021/09/08/10120583.html', 'https://jishou.rednet.cn/content/2021/09/14/10140988.html']
    # all_links = ['https://jishou.rednet.cn/content/2021/09/14/10140988.html']
    options = webdriver.ChromeOptions()

    options.add_argument('--headless')

    driver = webdriver.Chrome('/Users/mac/Downloads/chromedriver', options=options)
    driver.set_page_load_timeout(10)
    with open('crawled.txt', 'r') as f:
        crawled = [item.rstrip('\n') for item in f.readlines()]
        print('crawled', crawled)

    all_links = [l for l in all_links if l.startswith('http')]
    for link in all_links:
        if link in crawled:
            print('Found duplicate link: %s, skipping...' % link)
            continue
        try:
            driver.get(link)
            ps = driver.page_source
            html = etree.HTML(ps)
            title = html.xpath('//h1/text()')[0]

            if os.path.exists('output/%s.docx' % title):
                print('Found duplicate docx: %s, skipping...' % title)
                continue

            print('CRAWLING %s %s' % (title, link))
            driver.get_screenshot_as_file('intermediates/%s-start.png' % title)

            driver.execute_script("window.scrollTo(0, document.body.scrollHeight / 2)")
            time.sleep(0.2)
            driver.get_screenshot_as_file('intermediates/%s-middle.png' % title)

            driver.execute_script("window.scrollTo(0, document.body.scrollHeight)")
            time.sleep(0.2)
            driver.get_screenshot_as_file('intermediates/%s-end.png' % title)

            url = driver.current_url
            ws = WordSaver(html=html, title=title, url=url)
            ws.save()
            new_links_crawled.append(url)
        except:
            # raise e
            print(1111)
            traceback.print_exc()
            print('Error in crawling: %s' % link)
            with open('failed_to_crawl.txt', 'a') as f:
                f.write(link + '\n')
            continue

    with open('crawled.txt', 'a') as f:
        for item in new_links_crawled:
            f.write(item + '\n')
    driver.close()
