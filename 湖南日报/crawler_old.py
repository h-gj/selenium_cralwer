import os
import time
import traceback

import requests
from lxml import etree
from docx import Document
from docx.shared import Inches
from selenium import webdriver


class WordSaver:
    def __init__(self, html, title, url):
        self.html = html
        self.title = title
        self.url = url
        self.paragraphs = None
        self.parse()

    def parse(self):
        self.paragraphs = self.html.xpath('//div[@id="content"]//p/text()')
        if not self.paragraphs:
            print('getting content from another selector...')
            self.paragraphs = self.html.xpath('//div[@class="nei"]/ct/p//text()')
            if not self.paragraphs:
                print('failed again to find body: %s...' % self.url)

    def save(self):
        # 新建空白文档
        doc = Document()
        # 新增文档标题
        doc.add_heading(self.title, 0)
        doc.add_heading('原文链接 ' + self.url, 3)
        doc.add_picture('intermediates/%s-start.png' % self.title, width=Inches(5))
        doc.add_paragraph()
        doc.add_picture('intermediates/%s-middle.png' % self.title, width=Inches(5))
        doc.add_paragraph()
        doc.add_picture('intermediates/%s-end.png' % self.title, width=Inches(5))
        for p in self.paragraphs:
            # 创建段落描述
            doc.add_paragraph(p)
        doc.save('output/%s.docx' % self.title)


def get_page_item_links(kw, page=0, total_links=None):
    if not total_links:
        total_links = []
    url = 'http://so.voc.com.cn/cse/search?q=%s&p=%s&s=7639422230623402302&sti=1440' % (kw, page)
    content = requests.get(url).content.decode('utf-8')
    html = etree.HTML(content)
    links = html.xpath('//h3[@class="c-title"]/a/@href')
    if not links:
        return total_links

    if links[0] and links[1] in total_links:
        return total_links

    print('page links', page, links)
    total_links += links
    return get_page_item_links(kw, page=page + 1, total_links=total_links[:])


if __name__ == '__main__':
    keywords = ['望城县', '长沙市望城区', '望城区档案']
    all_links = []
    new_links_crawled = []
    for kw in keywords:
        all_links += get_page_item_links(kw)

    # all_links = ['http://www.chinaleifeng.com/content/2021/09/08/10120583.html', 'https://jishou.rednet.cn/content/2021/09/14/10140988.html']
    # all_links = ['https://jishou.rednet.cn/content/2021/09/14/10140988.html']
    options = webdriver.ChromeOptions()

    # options.add_argument('--headless')

    driver = webdriver.Chrome('/Users/mac/Downloads/chromedriver', options=options)
    driver.set_page_load_timeout(10)
    with open('crawled.txt', 'r') as f:
        crawled = [item.rstrip('\n') for item in f.readlines()]
        print('crawled', crawled)

    all_links = [l for l in all_links if l.startswith('http')]
    for link in all_links:
        if link in crawled:
            print('Found duplicate link: %s, skipping...' % link)
            continue
        try:
            driver.get(link)
            ps = driver.page_source
            html = etree.HTML(ps)
            title = html.xpath('//h1/text()')[0]

            if os.path.exists('output/%s.docx' % title):
                print('Found duplicate docx: %s, skipping...' % title)
                continue

            print('CRAWLING %s %s' % (title, link))
            driver.get_screenshot_as_file('intermediates/%s-start.png' % title)

            driver.execute_script("window.scrollTo(0, document.body.scrollHeight / 2)")
            time.sleep(0.2)
            driver.get_screenshot_as_file('intermediates/%s-middle.png' % title)

            driver.execute_script("window.scrollTo(0, document.body.scrollHeight)")
            time.sleep(0.2)
            driver.get_screenshot_as_file('intermediates/%s-end.png' % title)

            url = driver.current_url
            ws = WordSaver(html=html, title=title, url=url)
            ws.save()
            new_links_crawled.append(url)
        except Exception as e:
            traceback.print_exc()
            print('Error in crawling: %s' % link)
            with open('failed_to_crawl.txt', 'a') as f:
                f.write(link + '\n')
            continue

    with open('crawled.txt', 'a') as f:
        for item in new_links_crawled:
            f.write(item + '\n')
    driver.close()
